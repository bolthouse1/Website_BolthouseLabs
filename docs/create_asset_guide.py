"""
Generate the MyBodyPrism Media Asset Guide as a Word document.
Run: python docs/create_asset_guide.py
Output: docs/MyBodyPrism_Media_Asset_Guide.docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Heading styles
for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x0A, 0x0A, 0x12)

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()  # spacer
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('MyBodyPrism')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x00, 0xD4, 0xFF)
run.font.name = 'Calibri'
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Media Asset Guide')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
run.font.name = 'Calibri'

doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('www.bolthouselabs.com\nTeaser Website — Asset Placement Reference')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.font.name = 'Calibri'

doc.add_paragraph()
doc.add_paragraph()

instructions = doc.add_paragraph()
instructions.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = instructions.add_run(
    'HOW TO USE THIS DOCUMENT\n\n'
    'Each section below represents a placeholder on the live website.\n'
    'For each tagged slot, either:\n\n'
    '  Option A — Paste/embed the image or video screenshot directly into the grey box below each tag.\n'
    '  Option B — Provide files separately and reference the tag (e.g., "This file is MEDIA-2").\n\n'
    'When you return this document (or the files), Claude will drop each asset into the correct spot on the site.'
)
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
run.font.name = 'Calibri'

doc.add_page_break()

# ============================================================
# QUICK REFERENCE TABLE
# ============================================================
doc.add_heading('Quick Reference', level=1)

table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Tag', 'Type', 'Description', 'Aspect Ratio', 'Format']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(10)

assets = [
    ('MEDIA-1', 'Image', 'Clinical CT-PET slice — the "before" image', '16:9 or 4:3', 'PNG / JPG'),
    ('MEDIA-2', 'Video or Image', 'Hero 3D volumetric render — slow rotation, cinematic', '16:9', 'MP4 / PNG / WebP'),
    ('MEDIA-3', 'Image (optional)', 'Close-up cardiac detail render', '16:9', 'PNG / WebP'),
    ('MEDIA-4', 'Image or Video', 'MyBodyPrism viewer UI — loaded case', '16:9', 'MP4 / PNG / WebP'),
    ('MEDIA-5a/b/c', 'SVG icons (×3)', 'How It Works step icons (upgrade optional)', '1:1 (80×80)', 'SVG'),
]

for asset in assets:
    row = table.add_row()
    for i, val in enumerate(asset):
        cell = row.cells[i]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

# Set column widths
widths = [Inches(1.0), Inches(1.2), Inches(2.8), Inches(1.0), Inches(1.2)]
for row in table.rows:
    for i, width in enumerate(widths):
        row.cells[i].width = width

doc.add_paragraph()
doc.add_page_break()

# ============================================================
# HELPER: Add an asset section
# ============================================================
def add_asset_section(tag, section_name, description, specs, creative_direction, page_break=True):
    # Tag heading
    h = doc.add_heading(f'{tag}  —  {section_name}', level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x00, 0xA0, 0xCC)

    # Website section context
    p = doc.add_paragraph()
    run = p.add_run('Website Section: ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(section_name)
    run.font.size = Pt(11)

    # Description
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('What this asset shows:')
    run.bold = True
    run.font.size = Pt(11)
    p = doc.add_paragraph(description)
    p.style.font.size = Pt(11)

    # Specs
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Specifications:')
    run.bold = True
    run.font.size = Pt(11)
    for key, val in specs.items():
        p = doc.add_paragraph(f'  •  {key}: {val}')
        p.style.font.size = Pt(10)

    # Creative direction
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Creative Direction:')
    run.bold = True
    run.font.size = Pt(11)
    p = doc.add_paragraph(creative_direction)
    p.paragraph_format.space_after = Pt(12)

    # Drop zone box
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f'▼  PASTE / EMBED YOUR ASSET FOR {tag} BELOW  ▼')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x00, 0xA0, 0xCC)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Create a bordered "drop zone" table cell
    drop_table = doc.add_table(rows=1, cols=1)
    drop_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = drop_table.rows[0].cells[0]
    cell.width = Inches(6.5)

    # Set cell height via row height
    tr = drop_table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = trPr.makeelement(qn('w:trHeight'), {qn('w:val'): '4000', qn('w:hRule'): 'atLeast'})
    trPr.append(trHeight)

    # Add placeholder text inside the cell
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'\n\n\n[ Paste your {tag} image/screenshot here ]\n\n\n')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run.italic = True

    # Set cell shading to light grey
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F0F0F0'
    })
    shading.append(shading_elem)

    if page_break:
        doc.add_page_break()


# ============================================================
# ASSET SECTIONS
# ============================================================

add_asset_section(
    tag='MEDIA-1',
    section_name='The Diagnosis (Section 2)',
    description=(
        'A flat, clinical CT-PET slice image. This is the "before" — it should feel cold, '
        'clinical, and hard for a non-expert to interpret. This establishes the problem: '
        '"This is what patients get. Can you make sense of this?"'
    ),
    specs={
        'Aspect Ratio': '16:9 or 4:3',
        'Format': 'PNG or JPG',
        'Min Resolution': '1200px wide',
        'Background': 'Dark / black (to match site)',
    },
    creative_direction=(
        'Think of a standard radiology viewer — greyscale, maybe some color overlay from PET. '
        'A single slice or a small tiled view. Should look authentic and clinical. '
        'The viewer should NOT look modern or polished — the contrast with MEDIA-2 is the whole point. '
        'If possible, use a cardiac case (chest CT or PET-CT with cardiac focus).'
    ),
)

add_asset_section(
    tag='MEDIA-2',
    section_name='The Reveal — Hero Render (Section 4)',
    description=(
        'THIS IS THE SHOWSTOPPER. A stunning 3D volumetric render — the "after." '
        'Slow rotation, cinematic lighting, dark background. This is the moment the visitor '
        'goes "holy shit" and understands what MyBodyPrism does. The dramatic payoff of the '
        'entire narrative arc.'
    ),
    specs={
        'Aspect Ratio': '16:9',
        'Format': 'MP4 (preferred, looping video) or PNG/WebP (static)',
        'Min Resolution': '1920×1080 for video, 1600px wide for image',
        'Background': 'Dark / black (MUST blend with site)',
        'If video': 'Autoplay, muted, looping. 5-15 seconds ideal. Keep file < 10MB.',
    },
    creative_direction=(
        'Cinematic 3D body render — think medical visualization meets movie VFX. '
        'Could be a full torso, a heart, or a body segment rendered volumetrically. '
        'Dramatic lighting (rim light, subtle glow, depth). Slow rotation or camera drift. '
        'Color palette should lean cool (blues, cyans) with warm accent highlights. '
        'This should look like nothing a patient has ever seen from a hospital. '
        'DARK BACKGROUND IS CRITICAL — it must blend seamlessly into the page.'
    ),
)

add_asset_section(
    tag='MEDIA-3',
    section_name='The Reveal — Detail Shot (Section 4, Optional)',
    description=(
        'An optional close-up detail render. Could be a zoomed-in cardiac structure, '
        'vascular tree, or tissue detail. Adds depth to the reveal section if you have it.'
    ),
    specs={
        'Aspect Ratio': '16:9',
        'Format': 'PNG or WebP',
        'Min Resolution': '1200px wide',
        'Background': 'Dark / black',
    },
    creative_direction=(
        'Think macro photography but for 3D medical rendering. A close crop of something '
        'beautiful and intricate — coronary arteries, lung vasculature, cardiac chambers. '
        'Should feel intimate and detailed vs. the wide hero shot of MEDIA-2. '
        'Same cinematic lighting style. This is optional — skip if MEDIA-2 is strong enough on its own.'
    ),
)

add_asset_section(
    tag='MEDIA-4',
    section_name='Product Tease — UI Screenshot (Section 5)',
    description=(
        'A screenshot or screen recording of the MyBodyPrism application UI with a case loaded. '
        'This grounds the product — "this is real software, not just a concept." '
        'Should show the viewer interface with a 3D model loaded and controls visible.'
    ),
    specs={
        'Aspect Ratio': '16:9',
        'Format': 'PNG/WebP (screenshot) or MP4 (screen recording)',
        'Min Resolution': '1600px wide for image, 1920×1080 for video',
        'Background': 'App UI (dark theme preferred)',
        'If video': 'Short interaction — rotate, zoom, maybe toggle a view. 5-10 sec.',
    },
    creative_direction=(
        'Show the actual product UI looking polished and real. A 3D model should be loaded '
        'and visible. The UI should feel premium — dark theme, clean layout, professional controls. '
        'This is NOT the cinematic render — it\'s the product that creates the render. '
        'Think "here\'s the tool in your hands." '
        'If using a screen recording, show a smooth interaction (rotation, zoom) — no cursor fumbling.'
    ),
)

add_asset_section(
    tag='MEDIA-5a / 5b / 5c',
    section_name='How It Works — Step Icons (Section 6)',
    description=(
        'Three icons for the "How It Works" steps:\n'
        '  5a — "Load Your Scans" (upload/import concept)\n'
        '  5b — "We Transform It" (processing/magic concept)\n'
        '  5c — "Explore in 3D" (3D viewer/immersion concept)\n\n'
        'Basic SVG icons already exist on the site. This is an optional upgrade.'
    ),
    specs={
        'Size': '80×80px',
        'Format': 'SVG (preferred) or PNG with transparency',
        'Style': 'Monoline or minimal, cyan (#00d4ff) accent color',
    },
    creative_direction=(
        'Minimal, elegant line icons. Should feel premium and consistent as a set. '
        'Cyan stroke/accent on dark background. Think Apple-style simplicity. '
        'These are small — clarity over detail. Skip this if the current basic SVGs are fine for launch.'
    ),
    page_break=False,
)

# ============================================================
# FINAL NOTES PAGE
# ============================================================
doc.add_page_break()
doc.add_heading('Notes & Tips', level=1)

notes = [
    ('Dark backgrounds are mandatory',
     'The site is almost entirely dark (#050508). Any media with a white or light background will look jarring. '
     'Render on black, or we can add a vignette/overlay, but native dark is best.'),
    ('Video keeps file sizes small',
     'For MEDIA-2 and MEDIA-4, if using video, target under 10MB. H.264 MP4 is ideal. '
     'These will autoplay muted and loop — no audio needed.'),
    ('MEDIA-2 is the most important asset',
     'If you only have time/resources for one asset, make it MEDIA-2. The entire page narrative '
     'builds to that reveal moment. Everything else is supporting.'),
    ('"Good enough" is fine for launch',
     'Placeholder boxes work for now. Even one real asset (MEDIA-2) dramatically improves the page. '
     'You can always swap in better versions later — the tags make it easy.'),
    ('How to give assets to Claude',
     'Option A: Paste images directly into the grey boxes above, save, and share this doc.\n'
     'Option B: Send files separately and say "This is MEDIA-2" etc.\n'
     'Option C: Drop files in the project folder and point Claude at them.\n'
     'Any option works — the tags are what matter.'),
]

for title, body in notes:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}:  ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(body)
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(12)

# ============================================================
# SAVE
# ============================================================
output_path = r'c:\Website_BolthouseLabs\docs\MyBodyPrism_Media_Asset_Guide.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
